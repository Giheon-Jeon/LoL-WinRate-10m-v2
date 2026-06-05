"""report.md → report.docx 변환 스크립트"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

BODY_FONT = '맑은 고딕'
BODY_SIZE = 11
HEADING_SIZE = 13

def set_font(run, font_name=BODY_FONT, size=BODY_SIZE, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def set_line_spacing(paragraph, spacing=1.5):
    paragraph.paragraph_format.line_spacing = spacing

def add_page_number(paragraph):
    """페이지 번호 필드 추가"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_font(run, size=10)

def add_toc(document):
    """목차 자동 생성 필드"""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "여기를 우클릭 후 '필드 업데이트'를 선택하세요."
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    run = p.add_run(text)
    set_font(run, size=HEADING_SIZE, bold=True)
    set_line_spacing(p)
    return p

def add_paragraph(doc, text, bold=False, size=BODY_SIZE):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    set_line_spacing(p)
    return p

def parse_inline(text):
    """**bold**, *italic*, `code` 파싱하여 (text, bold, italic, code) 리스트 반환"""
    parts = []
    pattern = r'(\*\*[^*]+\*\*|`[^`]+`)'
    pieces = re.split(pattern, text)
    for piece in pieces:
        if not piece: continue
        if piece.startswith('**') and piece.endswith('**'):
            parts.append((piece[2:-2], True, False, False))
        elif piece.startswith('`') and piece.endswith('`'):
            parts.append((piece[1:-1], False, False, True))
        else:
            parts.append((piece, False, False, False))
    return parts

def add_styled_paragraph(doc, text, size=BODY_SIZE):
    p = doc.add_paragraph()
    for content, bold, italic, code in parse_inline(text):
        run = p.add_run(content)
        if code:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            set_font(run, size=size, bold=bold)
    set_line_spacing(p)
    return p

def add_table_from_markdown(doc, header, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(header))
    table.style = 'Light Grid Accent 1'
    table.autofit = True
    # 헤더
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        for content, bold, italic, code in parse_inline(cell_text.strip()):
            run = p.add_run(content)
            set_font(run, size=BODY_SIZE, bold=True)
    # 데이터 행
    for i, row in enumerate(rows, 1):
        for j, cell_text in enumerate(row):
            if j >= len(header): break
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            for content, bold, italic, code in parse_inline(cell_text.strip()):
                run = p.add_run(content)
                set_font(run, size=BODY_SIZE, bold=bold)
    return table

def parse_markdown_table(lines, start_idx):
    """마크다운 표 파싱. (header, rows, next_idx) 반환"""
    header_line = lines[start_idx].strip().strip('|')
    header = [c.strip() for c in header_line.split('|')]
    # 구분선 건너뛰기
    idx = start_idx + 2
    rows = []
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        row_line = lines[idx].strip().strip('|')
        row = [c.strip() for c in row_line.split('|')]
        rows.append(row)
        idx += 1
    return header, rows, idx

def is_table_start(lines, idx):
    if idx + 1 >= len(lines): return False
    if not lines[idx].strip().startswith('|'): return False
    next_line = lines[idx+1].strip()
    return bool(re.match(r'^\|[-:|\s]+\|$', next_line))

# ── 문서 생성 ──
doc = Document()

# 페이지 여백
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 표지 ──
for _ in range(5):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('League of Legends')
set_font(title_run, size=24, bold=True)
title_p.add_run().add_break()
subtitle_run = title_p.add_run('초반 15분 데이터 기반\n승패 예측 모델 개발')
set_font(subtitle_run, size=20, bold=True)

for _ in range(8):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run('팀  원 : ____________________________')
set_font(r, size=13)
info.add_run().add_break()
info.add_run().add_break()
r = info.add_run('제출일 : ____________________________')
set_font(r, size=13)

doc.add_page_break()

# ── 목차 ──
toc_heading = doc.add_paragraph()
toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = toc_heading.add_run('목  차')
set_font(r, size=18, bold=True)
doc.add_paragraph()
add_toc(doc)
doc.add_page_break()

# ── 페이지 번호 (하단 중앙) ──
footer = doc.sections[0].footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_number(footer_para)

# ── 본문 ──
with open('report.md', encoding='utf-8') as f:
    md = f.read()

lines = md.split('\n')
i = 0
in_code_block = False
code_buffer = []

while i < len(lines):
    line = lines[i]

    # 코드 블록 처리
    if line.strip().startswith('```'):
        if in_code_block:
            for code_line in code_buffer:
                p = doc.add_paragraph()
                run = p.add_run(code_line)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            code_buffer = []
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue
    if in_code_block:
        code_buffer.append(line)
        i += 1
        continue

    # 수평선
    if line.strip() == '---':
        i += 1
        continue

    # 헤더
    if line.startswith('# '):
        # 본문 시작 시 페이지 나누기 한번
        add_heading(doc, line[2:].strip(), level=1)
        i += 1
        continue
    if line.startswith('## '):
        add_heading(doc, line[3:].strip(), level=1)
        i += 1
        continue
    if line.startswith('### '):
        heading_text = line[4:].strip()
        add_heading(doc, heading_text, level=2)
        # 5.x 섹션에 그래프 placeholder 추가
        if re.match(r'^5\.\d+\s', heading_text):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run('[그래프 삽입 위치]')
            set_font(r, size=BODY_SIZE, bold=True, color=RGBColor(0x80, 0x80, 0x80))
            set_line_spacing(p)
        i += 1
        continue
    if line.startswith('#### '):
        add_heading(doc, line[5:].strip(), level=3)
        i += 1
        continue

    # 표 처리
    if is_table_start(lines, i):
        header, rows, next_idx = parse_markdown_table(lines, i)
        add_table_from_markdown(doc, header, rows)
        # 표 다음 빈 줄
        doc.add_paragraph()
        i = next_idx
        continue

    # 인용
    if line.startswith('> '):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        for content, bold, italic, code in parse_inline(line[2:]):
            run = p.add_run(content)
            set_font(run, size=BODY_SIZE, bold=bold)
            if bold:
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        set_line_spacing(p)
        i += 1
        continue

    # 리스트
    if re.match(r'^\s*[-*]\s+', line):
        # 들여쓰기 깊이 계산
        indent_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if indent_match:
            indent = len(indent_match.group(1)) // 2
            content = indent_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
            for c, b, it, code in parse_inline(content):
                run = p.add_run(c)
                set_font(run, size=BODY_SIZE, bold=b)
            set_line_spacing(p)
            i += 1
            continue

    # 번호 리스트
    if re.match(r'^\s*\d+\.\s+', line):
        m = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if m:
            indent = len(m.group(1)) // 2
            content = m.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
            for c, b, it, code in parse_inline(content):
                run = p.add_run(c)
                set_font(run, size=BODY_SIZE, bold=b)
            set_line_spacing(p)
            i += 1
            continue

    # 일반 단락
    if line.strip():
        add_styled_paragraph(doc, line)
        i += 1
        continue
    else:
        # 빈 줄
        i += 1
        continue

# 인사이트 분석 섹션 뒤에 그래프 자리 표시
# (이미 본문 처리 완료. 별도 placeholder 추가하려면 수동 삽입 필요)
# 본문에 그래프 placeholder 추가 — 인사이트 섹션 끝에 추가는 어려우니
# 본문 처리 중에 ### 5.x 헤더 뒤에 자동 삽입하는 게 좋겠음

# 대신 결론 직전에 일괄적으로 추가하는 방식은 의미 없음 — 본문 처리 시 5.x 섹션마다 추가하도록 후처리

doc.save('report.docx')
print(f"report.docx 저장 완료")
