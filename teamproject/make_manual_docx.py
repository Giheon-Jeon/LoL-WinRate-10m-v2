"""manual.docx 생성 스크립트 - 코드 실행 설명서"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = '맑은 고딕'
CODE_FONT = 'Courier New'

def set_font(run, font_name=BODY_FONT, size=11, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 12}
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    run = p.add_run(text)
    set_font(run, size=sizes.get(level, 11), bold=True)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bullet(doc, text, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
    p.paragraph_format.line_spacing = 1.5
    # **bold** 처리
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for piece in parts:
        if not piece: continue
        if piece.startswith('**') and piece.endswith('**'):
            run = p.add_run(piece[2:-2])
            set_font(run, bold=True)
        else:
            run = p.add_run(piece)
            set_font(run)
    return p

def add_code_block(doc, code_text):
    """Courier New 10pt 회색 배경 코드 블록"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.2
    # 회색 배경
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    # 테두리
    pBdr = OxmlElement('w:pBdr')
    for border_name in ('top','left','bottom','right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border)
    pPr.append(pBdr)

    run = p.add_run(code_text)
    run.font.name = CODE_FONT
    run.font.size = Pt(10)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), CODE_FONT)
    rFonts.set(qn('w:hAnsi'), CODE_FONT)
    return p

def add_table(doc, header, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(header))
    table.style = 'Light Grid Accent 1'
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            if j >= len(header): break
            cell = table.rows[i].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            set_font(run)
    return table

# ── 문서 생성 ──
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 표지 ──
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('LoL 승패 예측 프로젝트')
set_font(r, size=22, bold=True)
title_p.add_run().add_break()
r = title_p.add_run('코드 실행 설명서')
set_font(r, size=22, bold=True)

for _ in range(10):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run('팀  원 : ____________________________')
set_font(r, size=13)
info.add_run().add_break()
info.add_run().add_break()
r = info.add_run('제출일 : ____________________________')
set_font(r, size=13)

doc.add_page_break()

# ───────────────────────────────────────────────────────────────────
# 1. 프로젝트 구조
# ───────────────────────────────────────────────────────────────────
add_heading(doc, '1. 프로젝트 구조', level=1)

add_para(doc, '본 프로젝트는 LoL 솔로랭크 게임의 초반 15분 데이터를 활용한 승패 예측 모델을 구현하며, 다음과 같은 파일로 구성된다.')

add_heading(doc, '1.1 파일 구성', level=2)

add_table(doc,
    ['파일명', '구분', '역할'],
    [
        ['data.ipynb', '노트북', 'Riot API를 통한 데이터 수집 (8,410경기 / 80컬럼)'],
        ['preprocessing.ipynb', '노트북', '컬럼 제거 / 학습·테스트 분할 / 표준화'],
        ['modeling.ipynb', '노트북', 'LR, RF, XGBoost 학습 / 5-fold CV / Optuna 튜닝'],
        ['insight.ipynb', '노트북', '오브젝트 vs 골드차, 라인별 골드 효율 등 EDA'],
        ['insight_analysis.ipynb', '노트북', 'CS, KDA, Feature Importance 심층 분석'],
        ['lol_clean_final.csv', '데이터', '수집·정제된 최종 데이터셋'],
        ['xgb_tuned.pkl', '모델', 'Optuna 최적 XGBoost 모델 저장 파일'],
        ['report.md / report.docx', '문서', '프로젝트 최종 보고서'],
        ['manual.docx', '문서', '본 코드 실행 설명서'],
    ]
)

add_heading(doc, '1.2 실행 순서', level=2)
add_para(doc, '아래 순서대로 노트북을 실행하면 전체 파이프라인이 재현된다.')
add_bullet(doc, '**1단계** : data.ipynb — 데이터 수집 (선택사항, lol_clean_final.csv 제공 시 생략 가능)')
add_bullet(doc, '**2단계** : preprocessing.ipynb — 전처리 및 학습/테스트 분할 확인')
add_bullet(doc, '**3단계** : modeling.ipynb — 모델 학습, 교차검증, 하이퍼파라미터 튜닝')
add_bullet(doc, '**4단계** : insight.ipynb — 기본 EDA 인사이트 도출')
add_bullet(doc, '**5단계** : insight_analysis.ipynb — 심층 인사이트 분석')

doc.add_page_break()

# ───────────────────────────────────────────────────────────────────
# 2. 환경 설정
# ───────────────────────────────────────────────────────────────────
add_heading(doc, '2. 환경 설정', level=1)

add_heading(doc, '2.1 Python 버전', level=2)
add_para(doc, '권장: Python 3.10 이상')
add_para(doc, '본 프로젝트는 Python 3.10.0에서 개발 및 검증되었다.')

add_heading(doc, '2.2 필요 라이브러리', level=2)
add_para(doc, '아래는 본 프로젝트에서 사용하는 모든 외부 라이브러리이다.')

add_table(doc,
    ['라이브러리', '용도'],
    [
        ['pandas', '데이터 프레임 처리'],
        ['numpy', '수치 연산'],
        ['requests', 'Riot API HTTP 요청'],
        ['tqdm', '진행률 표시바'],
        ['matplotlib', '시각화 (그래프 작성)'],
        ['seaborn', '시각화 보조'],
        ['scikit-learn', '머신러닝 모델 (LR, RF, 전처리, 분할, CV)'],
        ['xgboost', 'XGBoost 분류 모델'],
        ['optuna', '하이퍼파라미터 탐색'],
        ['joblib', '모델 저장 / 로드'],
    ]
)

add_heading(doc, '2.3 설치 명령어', level=2)
add_para(doc, '터미널에서 아래 명령어를 실행하여 일괄 설치한다.')
add_code_block(doc,
    'pip install pandas numpy requests tqdm matplotlib seaborn scikit-learn xgboost optuna joblib'
)
add_para(doc, '개별 설치를 원하는 경우 다음과 같이 실행한다.')
add_code_block(doc,
    'pip install pandas\n'
    'pip install numpy\n'
    'pip install requests\n'
    'pip install tqdm\n'
    'pip install matplotlib\n'
    'pip install seaborn\n'
    'pip install scikit-learn\n'
    'pip install xgboost\n'
    'pip install optuna\n'
    'pip install joblib'
)

add_para(doc, '※ macOS 환경에서 한글 폰트(AppleGothic)를 사용하므로 Windows에서는 다음과 같이 변경 필요:')
add_code_block(doc,
    "plt.rcParams['font.family'] = 'Malgun Gothic'  # Windows\n"
    "# plt.rcParams['font.family'] = 'AppleGothic' # macOS"
)

doc.add_page_break()

# ───────────────────────────────────────────────────────────────────
# 3. 데이터 설정
# ───────────────────────────────────────────────────────────────────
add_heading(doc, '3. 데이터 설정', level=1)

add_heading(doc, '3.1 lol_clean_final.csv 파일 위치', level=2)
add_para(doc, '모든 노트북은 작업 디렉토리에 lol_clean_final.csv 파일이 존재한다고 가정한다. 노트북 파일들과 동일한 폴더에 데이터셋을 위치시킨다.')
add_code_block(doc,
    'teamproject/\n'
    '├── data.ipynb\n'
    '├── preprocessing.ipynb\n'
    '├── modeling.ipynb\n'
    '├── insight.ipynb\n'
    '├── insight_analysis.ipynb\n'
    '├── lol_clean_final.csv      ← 여기에 위치\n'
    '└── xgb_tuned.pkl'
)

add_heading(doc, '3.2 Riot API 키 발급', level=2)
add_para(doc, '데이터 수집(data.ipynb)을 직접 실행하려면 Riot API 키가 필요하다. 이미 제공된 lol_clean_final.csv를 사용한다면 이 단계는 생략 가능하다.')
add_bullet(doc, '**1단계** : Riot Developer Portal 접속 (https://developer.riotgames.com)')
add_bullet(doc, '**2단계** : Riot 계정으로 로그인')
add_bullet(doc, '**3단계** : "Personal API Key" 또는 "Development API Key" 발급')
add_bullet(doc, '**4단계** : 발급된 키 복사 (RGAPI-로 시작하는 문자열)')

add_para(doc, '※ Development API Key는 24시간마다 갱신이 필요하며, 분당 100건의 요청 제한이 있다.')

add_heading(doc, '3.3 API 키 코드에 삽입', level=2)
add_para(doc, 'data.ipynb의 첫 번째 셀에서 API_KEY 변수의 값을 발급받은 키로 교체한다.')
add_code_block(doc,
    'import requests\n'
    'import pandas as pd\n'
    'import time\n'
    'from tqdm import tqdm\n'
    '\n'
    'API_KEY  = "RGAPI-여기에-발급받은-키-입력"   # ← 이 부분\n'
    'HEADER   = {"X-Riot-Token": API_KEY}\n'
    'KR_URL   = "https://kr.api.riotgames.com"\n'
    'ASIA_URL = "https://asia.api.riotgames.com"'
)

doc.add_page_break()

# ───────────────────────────────────────────────────────────────────
# 4. 파일별 실행 방법
# ───────────────────────────────────────────────────────────────────
add_heading(doc, '4. 파일별 실행 방법', level=1)

# 4.1 data.ipynb
add_heading(doc, '4.1 data.ipynb', level=2)
add_para(doc, '목적 : Riot API를 통해 한국 서버 상위 티어 게임 데이터 수집 및 lol_clean_final.csv 생성')

add_para(doc, '주요 셀 :', bold=True)
add_table(doc,
    ['셀 번호', '내용'],
    [
        ['Cell 0', 'API 키, 엔드포인트 URL 설정'],
        ['Cell 1', 'DataDragon API로 챔피언 메타데이터 로드 (167개 챔피언)'],
        ['Cell 2', '팀 조합 분류 함수 정의 (classify_comp)'],
        ['Cell 3', '매치 ID 조회 함수 및 매치 전체 추출 함수 정의'],
        ['Cell 4~6', '테스트용 5게임 수집 (API 연결 확인)'],
        ['Cell 7', '본 수집 실행 (5개 티어 × 약 3,000 매치 ID 수집 → 8,410게임)'],
    ]
)

add_para(doc, '예상 소요 시간 : 약 5시간 47분 (요청 간 1.2초 sleep 포함)', bold=True)

add_para(doc, '주의사항 :', bold=True)
add_bullet(doc, 'API 키 만료 시 401 응답이 발생하므로 24시간 이내에 완료할 것')
add_bullet(doc, '500게임마다 lol_clean_{N}.csv로 자동 중간 저장 — 중단되어도 부분 데이터 보존')
add_bullet(doc, '약 27%의 매치는 조기 종료(15분 미만) 또는 API 오류로 실패함 (정상)')
add_bullet(doc, '이미 lol_clean_final.csv가 제공된 경우 이 노트북은 실행할 필요 없음')

# 4.2 preprocessing.ipynb
add_heading(doc, '4.2 preprocessing.ipynb', level=2)
add_para(doc, '목적 : 수집된 원본 데이터에서 모델링에 부적합한 컬럼 제거 및 학습/테스트 분할')

add_para(doc, '주요 셀 :', bold=True)
add_table(doc,
    ['셀 번호', '내용'],
    [
        ['Cell 0', 'CSV 로드 및 기본 정보 출력 (8,410게임 / 80컬럼 / 결측 0)'],
        ['Cell 1', '불필요 컬럼 25개 제거 → 최종 55개 컬럼'],
        ['Cell 2', 'train_test_split(test_size=0.2, stratify=y, random_state=42)'],
        ['Cell 3', 'StandardScaler로 표준화'],
    ]
)

add_para(doc, '예상 소요 시간 : 약 5초 이내', bold=True)

add_para(doc, '주의사항 :', bold=True)
add_bullet(doc, '이 노트북은 lol_clean_final.csv가 같은 폴더에 있어야 실행 가능')
add_bullet(doc, '모델링 단계에서 다시 동일한 전처리를 수행하므로 결과 저장은 하지 않음')

# 4.3 modeling.ipynb
add_heading(doc, '4.3 modeling.ipynb', level=2)
add_para(doc, '목적 : LR/RF/XGBoost 3종 모델 학습, 5-fold 교차검증, XGBoost Optuna 튜닝')

add_para(doc, '주요 셀 :', bold=True)
add_table(doc,
    ['셀 번호', '내용'],
    [
        ['Cell 0', '데이터 로드 및 전처리 (preprocessing.ipynb와 동일)'],
        ['Cell 1', 'V1 모델 학습 (팀 단위 9피처)'],
        ['Cell 2', 'V2 모델 학습 (라인 추가 54피처)'],
        ['Cell 3', 'V1 vs V2 성능 비교 시각화'],
        ['Cell 4', 'StratifiedKFold 5-fold 교차검증 (6개 모델)'],
        ['Cell 5', 'Optuna 50회 탐색으로 XGBoost 튜닝 → xgb_tuned.pkl 저장'],
    ]
)

add_para(doc, '예상 소요 시간 : 약 3~5분 (Optuna 튜닝 포함)', bold=True)

add_para(doc, '주의사항 :', bold=True)
add_bullet(doc, '교차검증 셀(Cell 4)에서 KeyError 발생 시 커널 재시작 후 재실행 권장')
add_bullet(doc, 'Optuna 셀(Cell 5)은 50회 trial을 수행하므로 다소 시간이 소요됨')
add_bullet(doc, 'optuna 미설치 시 자동으로 ModuleNotFoundError 발생 — pip install optuna 후 재실행')

# 4.4 insight.ipynb
add_heading(doc, '4.4 insight.ipynb', level=2)
add_para(doc, '목적 : 오브젝트 / 골드차 / 라인별 골드 효율 / 라인 붕괴 시 회복 가능성 분석')

add_para(doc, '주요 셀 :', bold=True)
add_table(doc,
    ['셀 번호', '내용'],
    [
        ['Cell 0', '오브젝트 vs 골드차 LR 계수 비교 (골드차 계수 2.00으로 압도)'],
        ['Cell 1', '골드차 구간별 블루팀 승률 (3000↑ 시 91%)'],
        ['Cell 2', '라인별 골드 효율 분석 (정글 상관계수 0.394로 1위)'],
        ['Cell 3', '골드 하위 25% / 데스 상위 25% 시 라인별 승률'],
    ]
)

add_para(doc, '예상 소요 시간 : 약 10초 이내', bold=True)

# 4.5 insight_analysis.ipynb
add_heading(doc, '4.5 insight_analysis.ipynb', level=2)
add_para(doc, '목적 : CS, KDA, Feature Importance 등 심층 인사이트 분석')

add_para(doc, '주요 셀 :', bold=True)
add_table(doc,
    ['셀 번호', '내용'],
    [
        ['Cell 0', '공통 import, 데이터 로드, 라인 한글 매핑'],
        ['Cell 1', '분석1: CS vs 킬 상관계수 비교 (모든 라인에서 킬이 더 영향)'],
        ['Cell 2', '분석2: CS 격차 → 승률 (원딜 CS 격차 ±20에서 승률 41%p 변동)'],
        ['Cell 3', '분석3: 라인별 KDA 효율 (정글 KDA 1위)'],
        ['Cell 4', '분석4: XGBoost Feature Importance 상위 15개'],
    ]
)

add_para(doc, '예상 소요 시간 : 약 30초 이내 (Cell 4의 XGBoost 학습 포함)', bold=True)

doc.add_page_break()

# ───────────────────────────────────────────────────────────────────
# 5. 출력 파일 설명
# ───────────────────────────────────────────────────────────────────
add_heading(doc, '5. 출력 파일 설명', level=1)
add_para(doc, '코드 실행 시 생성되는 파일 목록과 용도는 다음과 같다.')

add_table(doc,
    ['파일명', '생성 노트북', '용도'],
    [
        ['lol_clean_final.csv', 'data.ipynb', '수집 완료된 최종 데이터셋 (8,410경기 × 80컬럼)'],
        ['lol_clean_{N}.csv', 'data.ipynb', '수집 중간 저장 파일 (500게임마다, N=500/1000/...)'],
        ['lol_test_5games.csv', 'data.ipynb', '테스트 수집 결과 (5게임만 수집, API 연결 확인용)'],
        ['xgb_tuned.pkl', 'modeling.ipynb', 'Optuna 최적 XGBoost 모델 (joblib 직렬화)'],
    ]
)

add_para(doc, 'xgb_tuned.pkl 로드 방법 :', bold=True)
add_code_block(doc,
    'import joblib\n'
    'import pandas as pd\n'
    '\n'
    'model = joblib.load("xgb_tuned.pkl")\n'
    'df = pd.read_csv("lol_clean_final.csv")\n'
    '# (전처리 동일하게 수행)\n'
    'pred = model.predict(X_new)\n'
    'proba = model.predict_proba(X_new)[:, 1]'
)

doc.add_page_break()

# ───────────────────────────────────────────────────────────────────
# 6. 자주 발생하는 오류
# ───────────────────────────────────────────────────────────────────
add_heading(doc, '6. 자주 발생하는 오류', level=1)

# 6.1
add_heading(doc, '6.1 Riot API rate limit (429 Too Many Requests)', level=2)
add_para(doc, '증상 : data.ipynb 실행 중 status_code 429 응답이 반복되며 수집 실패')
add_para(doc, '원인 : Riot Development API의 분당 100건 요청 한도 초과')
add_para(doc, '해결 :', bold=True)
add_bullet(doc, 'time.sleep(1.2) 값을 1.5~2.0으로 늘려 요청 간격 확장')
add_bullet(doc, '잠시 대기 후 (1~2분) 재실행')
add_bullet(doc, 'Production API Key 신청 (Riot Games 승인 필요)')

# 6.2
add_heading(doc, '6.2 API 키 만료 (401 Unauthorized)', level=2)
add_para(doc, '증상 : 401 Unauthorized 응답으로 모든 매치 수집 실패')
add_para(doc, '원인 : Development API Key는 24시간 만료')
add_para(doc, '해결 :', bold=True)
add_bullet(doc, 'Riot Developer Portal에서 새 키 발급 후 코드의 API_KEY 변수 교체')

# 6.3
add_heading(doc, '6.3 ModuleNotFoundError: No module named "xgboost"', level=2)
add_para(doc, '증상 : XGBoost 또는 다른 라이브러리 미설치 오류')
add_para(doc, '해결 :', bold=True)
add_code_block(doc, 'pip install xgboost  # 해당 라이브러리명으로 변경')

# 6.4
add_heading(doc, '6.4 KeyError: "ROCAUC" (modeling.ipynb 교차검증 셀)', level=2)
add_para(doc, '증상 : 박스플롯 시각화 단계에서 KeyError 발생')
add_para(doc, '원인 : Jupyter가 메모리에 구버전 셀을 캐시한 상태')
add_para(doc, '해결 :', bold=True)
add_bullet(doc, 'Kernel > Restart Kernel and Run All Cells 실행')
add_bullet(doc, '또는 해당 셀에서 metric_name.replace(...) 라인을 직접 삭제 후 재실행')

# 6.5
add_heading(doc, '6.5 한글 폰트 깨짐 (□□□ 표시)', level=2)
add_para(doc, '증상 : 그래프의 한글 라벨이 모두 사각형으로 표시')
add_para(doc, '원인 : OS별 한글 폰트 차이')
add_para(doc, '해결 :', bold=True)
add_code_block(doc,
    "# macOS\n"
    "plt.rcParams['font.family'] = 'AppleGothic'\n"
    "\n"
    "# Windows\n"
    "plt.rcParams['font.family'] = 'Malgun Gothic'\n"
    "\n"
    "# Linux (Ubuntu)\n"
    "plt.rcParams['font.family'] = 'NanumGothic'\n"
    "\n"
    "# 마이너스 부호 깨짐 방지\n"
    "plt.rcParams['axes.unicode_minus'] = False"
)

# 6.6
add_heading(doc, '6.6 lol_clean_final.csv FileNotFoundError', level=2)
add_para(doc, '증상 : preprocessing.ipynb / modeling.ipynb 등 실행 시 CSV 파일을 찾지 못함')
add_para(doc, '원인 : 노트북의 작업 디렉토리와 CSV 위치 불일치')
add_para(doc, '해결 :', bold=True)
add_bullet(doc, '노트북과 같은 폴더에 lol_clean_final.csv 위치 확인')
add_bullet(doc, '또는 절대 경로로 변경: pd.read_csv("/full/path/to/lol_clean_final.csv")')

# 6.7
add_heading(doc, '6.7 Optuna 진행률 바 표시 오류', level=2)
add_para(doc, '증상 : show_progress_bar=True 사용 시 환경에 따라 오류')
add_para(doc, '해결 :', bold=True)
add_code_block(doc,
    'study.optimize(objective, n_trials=50)  # show_progress_bar 제거'
)

# 6.8
add_heading(doc, '6.8 XGBoost / scikit-learn 버전 충돌', level=2)
add_para(doc, '증상 : eval_metric, use_label_encoder 등 파라미터 미지원 경고')
add_para(doc, '권장 버전 :', bold=True)
add_bullet(doc, 'xgboost ≥ 2.0')
add_bullet(doc, 'scikit-learn ≥ 1.3')
add_bullet(doc, 'optuna ≥ 3.0')
add_code_block(doc,
    'pip install --upgrade xgboost scikit-learn optuna'
)

doc.save('manual.docx')
print("manual.docx 저장 완료")
